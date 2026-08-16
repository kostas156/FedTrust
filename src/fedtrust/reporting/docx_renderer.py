"""DOCX renderer for professional FedTrust assessment reports."""

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from fedtrust.reporting.models import AssessmentReport, AssessmentSection, AssessmentSeverity


class DocxRenderer:
    """Render a FedTrust assessment report as a DOCX document."""

    def render(
        self,
        report: AssessmentReport,
        output_path: Path,
        chart_paths: dict[str, Path] | None = None,
    ) -> Path:
        """Render an assessment report to a DOCX file."""
        document = Document()

        self._configure_document(document)
        self._configure_styles(document)
        self._add_cover_page(document, report)
        self._add_executive_summary(document, report)
        self._add_overall_assessment(document, report)

        for section in report.sections:
            chart_path = None

            if chart_paths is not None:
                chart_path = chart_paths.get(section.name)

            self._add_section(document, section, chart_path)

        self._add_metadata(document, report)
        self._add_footer(document)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))

        return output_path

    @staticmethod
    def _configure_document(document: DocumentObject) -> None:
        """Configure page size and margins."""
        for section in document.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

    @staticmethod
    def _configure_styles(document: DocumentObject) -> None:
        """Configure the document's base typography."""
        styles = document.styles

        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)

        title = styles["Title"]
        title.font.name = "Aptos Display"
        title.font.size = Pt(26)
        title.font.bold = True

        heading_1 = styles["Heading 1"]
        heading_1.font.name = "Aptos Display"
        heading_1.font.size = Pt(18)
        heading_1.font.bold = True
        heading_1.font.color.rgb = RGBColor(31, 78, 121)

        heading_2 = styles["Heading 2"]
        heading_2.font.name = "Aptos Display"
        heading_2.font.size = Pt(13)
        heading_2.font.bold = True
        heading_2.font.color.rgb = RGBColor(68, 68, 68)

    @staticmethod
    def _add_cover_page(document: DocumentObject, report: AssessmentReport) -> None:
        """Add the report cover page."""
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph.paragraph_format.space_before = Pt(90)

        run = paragraph.add_run("FEDTRUST")
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 78, 121)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(12)

        run = title.add_run(report.title)
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(26)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_before = Pt(10)

        run = subtitle.add_run("Trustworthiness Assessment for Federated Learning Systems")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(90, 90, 90)

        severity = document.add_paragraph()
        severity.alignment = WD_ALIGN_PARAGRAPH.CENTER
        severity.paragraph_format.space_before = Pt(35)

        run = severity.add_run(f"OVERALL ASSESSMENT: {report.overall_severity.value.upper()}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = DocxRenderer._severity_color(report.overall_severity)

        document.add_page_break()

    @staticmethod
    def _add_executive_summary(document: DocumentObject, report: AssessmentReport) -> None:
        """Add the executive summary section."""
        document.add_heading("Executive Summary", level=1)

        paragraph = document.add_paragraph(report.executive_summary)
        paragraph.paragraph_format.space_after = Pt(12)

    @staticmethod
    def _add_overall_assessment(document: DocumentObject, report: AssessmentReport) -> None:
        """Add the overall assessment section."""
        document.add_heading("Overall Assessment", level=1)

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        cells = table.rows[0].cells
        cells[0].text = "Severity"
        cells[1].text = report.overall_severity.value.upper()

        DocxRenderer._style_header_cell(cells[0])

        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[1].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = DocxRenderer._severity_color(report.overall_severity)

    @staticmethod
    def _add_section(
        document: DocumentObject, section: AssessmentSection, chart_path: Path | None
    ) -> None:
        """Add one assessment section."""
        document.add_heading(DocxRenderer._display_section_name(section.name), level=1)

        summary = document.add_paragraph(section.summary)
        summary.paragraph_format.space_after = Pt(10)

        if section.metrics:
            DocxRenderer._add_metrics(document, section)

        if section.findings:
            DocxRenderer._add_findings(document, section)

        if section.recommendations:
            DocxRenderer._add_recommendations(document, section)

        if chart_path is not None and chart_path.exists():
            DocxRenderer._add_chart(
                document,
                section,
                chart_path,
            )

    @staticmethod
    def _add_metrics(document: DocumentObject, section: AssessmentSection) -> None:
        """Add the metrics table."""
        document.add_heading("Metrics", level=2)

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        headers = table.rows[0].cells
        headers[0].text = "Metric"
        headers[1].text = "Value"
        headers[2].text = "Unit"

        for cell in headers:
            DocxRenderer._style_header_cell(cell)

        for metric in section.metrics:
            cells = table.add_row().cells
            cells[0].text = DocxRenderer._display_metric_name(str(metric["name"]))
            cells[1].text = DocxRenderer._format_metric_value(metric["value"])
            cells[2].text = DocxRenderer._display_unit(metric.get("unit"))

    @staticmethod
    def _add_findings(
        document: DocumentObject,
        section: AssessmentSection,
    ) -> None:
        """Add findings and their evidence."""
        document.add_heading("Findings", level=2)

        for finding in section.findings:
            paragraph = document.add_paragraph(style="List Bullet")

            title_run = paragraph.add_run(f"{finding.title} [{finding.severity.value.upper()}]")
            title_run.bold = True
            title_run.font.color.rgb = DocxRenderer._severity_color(finding.severity)

            paragraph.add_run(f"\n{finding.description}")

            for evidence in finding.evidence:
                evidence_paragraph = document.add_paragraph(style="List Bullet 2")
                evidence_run = evidence_paragraph.add_run(f"Evidence: {evidence}")
                evidence_run.italic = True
                evidence_run.font.color.rgb = RGBColor(90, 90, 90)

    @staticmethod
    def _add_recommendations(
        document: DocumentObject,
        section: AssessmentSection,
    ) -> None:
        """Add recommendations."""
        document.add_heading("Recommendations", level=2)

        for recommendation in section.recommendations:
            paragraph = document.add_paragraph(style="List Bullet")

            title_run = paragraph.add_run(
                f"{recommendation.title} [{recommendation.priority.value.upper()}]"
            )
            title_run.bold = True
            title_run.font.color.rgb = DocxRenderer._severity_color(recommendation.priority)

            paragraph.add_run(f"\n{recommendation.description}")

    @staticmethod
    def _add_chart(
        document: DocumentObject,
        section: AssessmentSection,
        chart_path: Path,
    ) -> None:
        """Add an embedded chart and caption."""
        interpretation = document.add_paragraph()

        interpretation_run = interpretation.add_run(
            "Interpretation: The ROC curve illustrates how effectively "
            "the membership inference attack separates members from "
            "non-members. A curve closer to the upper-left region "
            "indicates stronger attack discrimination."
        )

        interpretation_run.font.size = Pt(9)
        interpretation_run.font.color.rgb = RGBColor(90, 90, 90)

        document.add_heading("Visual Evidence", level=2)

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph.add_run().add_picture(
            str(chart_path),
            width=Inches(6.0),
        )

        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = caption.add_run(
            f"Figure — {DocxRenderer._display_section_name(section.name)} evaluation evidence"
        )
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)

    @staticmethod
    def _add_metadata(
        document: DocumentObject,
        report: AssessmentReport,
    ) -> None:
        """Add report metadata as a technical appendix."""
        if not report.metadata:
            return

        document.add_page_break()
        document.add_heading("Technical Metadata", level=1)

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        headers = table.rows[0].cells
        headers[0].text = "Property"
        headers[1].text = "Value"

        for cell in headers:
            DocxRenderer._style_header_cell(cell)

        for key, value in report.metadata.items():
            cells = table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = str(value)

    @staticmethod
    def _add_footer(document: DocumentObject) -> None:
        """Add footer text and page numbers."""
        for section in document.sections:
            footer = section.footer

            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run("FedTrust • Trustworthiness Assessment • ")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(120, 120, 120)

            paragraph.add_run("Page ")
            field = OxmlElement("w:fldSimple")
            field.set(qn("w:instr"), "PAGE")
            paragraph._p.append(field)

    @staticmethod
    def _style_header_cell(cell: Any) -> None:
        """Apply consistent styling to a table header cell."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shading)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    @staticmethod
    def _severity_color(
        severity: AssessmentSeverity,
    ) -> RGBColor:
        """Return the display color associated with a severity."""
        colors = {
            AssessmentSeverity.INFO: RGBColor(90, 90, 90),
            AssessmentSeverity.LOW: RGBColor(84, 130, 53),
            AssessmentSeverity.MEDIUM: RGBColor(191, 144, 0),
            AssessmentSeverity.HIGH: RGBColor(192, 0, 0),
            AssessmentSeverity.CRITICAL: RGBColor(128, 0, 0),
        }

        return colors[severity]

    @staticmethod
    def _format_metric_value(value: Any) -> str:
        """Format metric values for human-readable presentation."""
        if isinstance(value, float):
            return f"{value:.3f}"

        return str(value)

    @staticmethod
    def _display_section_name(name: str) -> str:
        """Convert an internal section name to a human-readable label."""
        labels = {
            "classification": "Classification Performance",
            "membership_inference": "Membership Inference Privacy",
        }

        return labels.get(
            name,
            name.replace("_", " ").title(),
        )

    @staticmethod
    def _display_metric_name(name: str) -> str:
        """Convert an internal metric name to a human-readable label."""
        labels = {
            "accuracy": "Accuracy",
            "mia_auc": "MIA ROC-AUC",
        }

        return labels.get(
            name,
            name.replace("_", " ").title(),
        )

    @staticmethod
    def _display_unit(unit: str | None) -> str:
        """Convert an internal metric unit to a human-readable label."""
        labels = {
            "ratio": "Ratio",
            "roc_auc": "ROC-AUC",
        }

        if unit is None:
            return ""

        return labels.get(
            unit,
            unit.replace("_", " ").upper(),
        )
